import json
import os
import re
from collections import defaultdict
from datetime import datetime
from html import unescape

import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

target_date = datetime(2024, 5, 2)


def latex_to_human_readable(latex_str):
    # Remove $...$ delimiters
    latex_str = re.sub(r"\$(.*?)\$", r"\1", latex_str)

    simple_latex_to_text = {
        "\\ll": "<<",
        "\\alpha": "alpha",
        "\\epsilon": "epsilon",
        "\\widetilde": "widetilde",
        "\\in": "in",
        "\\leq": "<=",
        "\\geq": ">=",
        "\\pm": "±",
        "\\times": "x",
        "\\sim": "~",
        "\\approx": "≈",
        "\\neq": "≠",
        "\\cdot": "·",
        "\\ldots": "...",
        "\\cdots": "...",
        "\\vdots": "...",
        "\\ddots": "...",
        "\\forall": "for all",
        "\\exists": "exists",
        "\\nabla": "nabla",
        "\\partial": "partial",
        "\\{": "{",
        "\\}": "}",
        "\\:": " ",  # Small space
        "\\,": " ",  # Thin space
        "\\;": " ",  # Thick space
        "\\!": "",  # Negative space
        "_": "_",  # Subscript
    }

    for latex, text in simple_latex_to_text.items():
        latex_str = latex_str.replace(latex, text)

    single_arg_pattern = re.compile(r"\\(\w+){(.*?)}")
    latex_str = single_arg_pattern.sub(r"\2", latex_str)

    latex_str = latex_str.replace("``", '"').replace("''", '"')

    latex_str = latex_str.replace("--", "–")

    return unescape(latex_str)


def add_hyperlink_title(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)
    new_run = docx.oxml.shared.OxmlElement("w:r")
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    A function to add a hyperlink to a paragraph.
    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: The URL of the hyperlink.
    :param text: The text to be displayed for the link.
    :param color: The color of the hyperlink, in RGB format. Defaults to blue.
    :param underline: Whether the hyperlink should be underlined. Defaults to True.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)

    rPr = OxmlElement("w:rPr")
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    new_run.append(rPr)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def main():
    keywords = ["Gaussian Splatting", "Neural Radiance", "NeRF"]
    document = Document()
    p = document.add_paragraph()
    keyword_line = ", ".join(keywords)
    run = p.add_run(
        f"Here is your personalized list of papers from arXiv, created using the keywords {keyword_line}.\n\nThis list is courtesy of AtomikLabs, an open science project by "
    )
    run.font.color.rg = RGBColor(0, 0, 0)
    add_hyperlink(p, "https://www.linkedin.com/in/bradley-edwards-dev/", "Brad Edwards", underline=True)
    run = p.add_run(", a software and machine learning engineer from Vancouver, BC.")
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = document.add_paragraph()
    run = p.add_run("Thank you to arXiv for use of its open access interoperability.")
    run.font.color.rgb = RGBColor(0, 0, 0)
    document.add_heading("Summaries", level=2)
    with open("../../services/arxiv_summaries/src/records.json") as f:
        data = json.load(f)
        found_records = defaultdict(lambda: [])

        for record in data:
            for keyword in keywords:
                if (
                    datetime.strptime(record["date"], "%Y-%m-%d") >= target_date
                    and keyword.lower() in record["abstract"].lower()
                ):
                    found_records[keyword].append(record)
                    break

        for keyword in keywords:
            if not found_records[keyword]:
                continue
            document.add_heading(keyword, level=3)
            for record in found_records[keyword]:
                title_pdf_paragraph = document.add_paragraph()
                cleaned_title = re.sub(r"\n\s*", " ", record["title"])
                add_hyperlink_title(title_pdf_paragraph, cleaned_title, record["abstract_url"])
                title_pdf_paragraph.add_run(" [")
                add_hyperlink_title(title_pdf_paragraph, "PDF", record["abstract_url"].replace("/abs/", "/pdf/"))
                title_pdf_paragraph.add_run("]")
                authors = [f"{author['first_name']} {author['last_name']}" for author in record["authors"]]
                document.add_paragraph("by " + ", ".join(authors))
                document.add_paragraph(record["date"])
                cleaned_paragraph = re.sub(r"\n\s*", " ", record["abstract"])
                no_latex_paragraph = latex_to_human_readable(cleaned_paragraph)
                document.add_paragraph(no_latex_paragraph)

    try:
        os.remove("custom_newsletter.docx")
    except Exception as e:
        print(e)
    document.save("custom_newsletter.docx")


if __name__ == "__main__":
    main()
