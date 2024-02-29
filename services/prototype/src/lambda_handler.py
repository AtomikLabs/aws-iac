import json
import logging
import os
import re
import time
from collections import defaultdict
from datetime import date, datetime, timedelta
from html import unescape
from typing import List

import boto3
import defusedxml.ElementTree as ET
import docx
import requests
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI

from .storage_manager import StorageManager

logger = logging.getLogger(__name__)
logging.getLogger().setLevel(logging.INFO)
TEST = False
AURORA_CLUSTER_ARN = ""
BASE_URL = ""
BUCKET_NAME = ""
DB_CREDENTIALS_SECRET_ARN = ""  # nosec
DATABASE = ""
SUMMARY_SET = ""
OPENAI_KEY = ""

cs_categories_inverted = {
    "Computer Science - Artifical Intelligence": "AI",
    "Computer Science - Hardware Architecture": "AR",
    "Computer Science - Computational Complexity": "CC",
    "Computer Science - Computational Engineering, Finance, and Science": "CE",
    "Computer Science - Computational Geometry": "CG",
    "Computer Science - Computation and Language": "CL",
    "Computer Science - Cryptography and Security": "CR",
    "Computer Science - Computer Vision and Pattern Recognition": "CV",
    "Computer Science - Computers and Society": "CY",
    "Computer Science - Databases": "DB",
    "Computer Science - Distributed, Parallel, and Cluster Computing": "DC",
    "Computer Science - Digital Libraries": "DL",
    "Computer Science - Discrete Mathematics": "DM",
    "Computer Science - Data Structures and Algorithms": "DS",
    "Computer Science - Emerging Technologies": "ET",
    "Computer Science - Formal Languages and Automata Theory": "FL",
    "Computer Science - General Literature": "GL",
    "Computer Science - Graphics": "GR",
    "Computer Science - Computer Science and Game Theory": "GT",
    "Computer Science - Human-Computer Interaction": "HC",
    "Computer Science - Information Retrieval": "IR",
    "Computer Science - Information Theory": "IT",
    "Computer Science - Machine Learning": "LG",
    "Computer Science - Logic in Computer Science": "LO",
    "Computer Science - Multiagent Systems": "MA",
    "Computer Science - Multimedia": "MM",
    "Computer Science - Mathematical Software": "MS",
    "Computer Science - Numerical Analysis": "NA",
    "Computer Science - Neural and Evolutionary Computing": "NE",
    "Computer Science - Networking and Internet Architecture": "NI",
    "Computer Science - Other Computer Science": "OH",
    "Computer Science - Operating Systems": "OS",
    "Computer Science - Performance": "PF",
    "Computer Science - Programming Languages": "PL",
    "Computer Science - Robotics": "RO",
    "Computer Science - Symbolic Computation": "SC",
    "Computer Science - Sound": "SD",
    "Computer Science - Software Engineering": "SE",
    "Computer Science - Social and Information Networks": "SI",
    "Computer Science - Systems and Control": "SY",
}


def fetch_data(base_url: str, from_date: str, set: str) -> list:
    full_xml_responses = []
    params = {"verb": "ListRecords", "set": set, "metadataPrefix": "oai_dc", "from": from_date}

    backoff_times = [30, 120]

    while True:
        try:
            logging.info(f"Fetching data with parameters: {params}")
            response = requests.get(base_url, params=params, timeout=20)
            response.raise_for_status()
            full_xml_responses.append(response.text)
            root = ET.fromstring(response.content)
            resumption_token_element = root.find(".//{http://www.openarchives.org/OAI/2.0/}resumptionToken")

            if resumption_token_element is not None and resumption_token_element.text:
                logging.info(f"Found resumptionToken: {resumption_token_element.text}")
                print(f"Found resumptionToken: {resumption_token_element.text}")
                time.sleep(5)
                params = {"verb": "ListRecords", "resumptionToken": resumption_token_element.text}
            else:
                break

        except requests.exceptions.HTTPError as e:
            logging.error(f"HTTP error occurred: {e}")
            logging.error(f"Response content: {response.text}")
            print(e)

            if response.status_code == 503:
                backoff_time = response.headers.get("Retry-After", backoff_times.pop(0) if backoff_times else 30)
                logging.warning(f"Received 503 error, backing off for {backoff_time} seconds.")
                print(f"Received 503 error, backing off for {backoff_time} seconds.")
                time.sleep(int(backoff_time))
                continue

            break

        except Exception as e:
            logging.error(f"An unexpected error occurred: {e}")
            break

    return full_xml_responses


def parse_xml_data(xml_data: str, from_date: str) -> dict:
    extracted_data_chunk = defaultdict(list)

    try:
        root = ET.fromstring(xml_data)
        ns = {"oai": "http://www.openarchives.org/OAI/2.0/", "dc": "http://purl.org/dc/elements/1.1/"}

        for record in root.findall(".//oai:record", ns):
            date_elements = record.findall(".//dc:date", ns)
            if len(date_elements) != 1:
                continue

            identifier = record.find(".//oai:identifier", ns).text
            abstract_url = record.find(".//dc:identifier", ns).text

            creators_elements = record.findall(".//dc:creator", ns)
            authors = []
            for creator in creators_elements:
                name_parts = creator.text.split(", ", 1)
                last_name = name_parts[0]
                first_name = name_parts[1] if len(name_parts) > 1 else ""
                authors.append({"last_name": last_name, "first_name": first_name})

            # Find all subjects
            subjects_elements = record.findall(".//dc:subject", ns)
            categories = [cs_categories_inverted.get(subject.text, "") for subject in subjects_elements]
            # Remove empty strings
            categories = list(filter(None, categories))
            primary_category = categories[0] if categories else ""

            abstract = record.find(".//dc:description", ns).text.replace("\n", " ")
            title = record.find(".//dc:title", ns).text.replace("\n", "")
            date = date_elements[0].text
            group = "cs"

            extracted_data_chunk["records"].append(
                {
                    "identifier": identifier,
                    "abstract_url": abstract_url,
                    "authors": authors,
                    "primary_category": primary_category,
                    "categories": categories,  # All categories
                    "abstract": abstract,
                    "title": title,
                    "date": date,
                    "group": group,
                }
            )

    except ET.ParseError as e:
        print(f"Parse error: {e}")

    return extracted_data_chunk


def write_to_files(data: list, file_paths: dict) -> None:
    """Writes the extracted fields to different files.

    Args:
        data (list): A list of dictionaries containing the extracted fields.
        file_paths (dict): A dictionary with field names as keys and file paths as values.

    Returns:
        None
    """
    try:
        # Aggregating all 'records' from each dictionary in the list
        aggregated_records = [record for d in data for record in d.get("records", [])]

        with open(file_paths["records"], "w") as f:
            json_data = json.dumps(aggregated_records, indent=4)
            f.write(json_data)

        logging.info(f"Successfully wrote records data to {file_paths['records']}.")

    except Exception as e:
        logging.error(f"Failed to write data: {e}")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)
    new_run = docx.oxml.shared.OxmlElement("w:r")
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def latex_to_human_readable(latex_str):
    # Remove $...$ delimiters
    latex_str = re.sub(r"\$(.*?)\$", r"\1", latex_str)

    simple_latex_to_text = {
        "\\ll": "<<",
        "\\alpha": "alpha",
        "\\epsilon": "epsilon",
        "\\widetilde": "widetilde",
        "\\in": "in",
        "\\leq": "<=",
        "\\geq": ">=",
        "\\pm": "±",
        "\\times": "x",
        "\\sim": "~",
        "\\approx": "≈",
        "\\neq": "≠",
        "\\cdot": "·",
        "\\ldots": "...",
        "\\cdots": "...",
        "\\vdots": "...",
        "\\ddots": "...",
        "\\forall": "for all",
        "\\exists": "exists",
        "\\nabla": "nabla",
        "\\partial": "partial",
        "\\{": "{",
        "\\}": "}",
        "\\:": " ",  # Small space
        "\\,": " ",  # Thin space
        "\\;": " ",  # Thick space
        "\\!": "",  # Negative space
        "_": "_",  # Subscript
    }

    for latex, text in simple_latex_to_text.items():
        latex_str = latex_str.replace(latex, text)

    single_arg_pattern = re.compile(r"\\(\w+){(.*?)}")
    latex_str = single_arg_pattern.sub(r"\2", latex_str)

    latex_str = latex_str.replace("``", '"').replace("''", '"')

    latex_str = latex_str.replace("--", "–")

    return unescape(latex_str)


def create_full_show_notes(categories: list, records: list, research_date: str, group: str, themes: str):
    for category in categories:
        doc = Document()
        intro = ""
        theme_header = "Research Themes (AI-Generated)"
        thank_you = "Thank you to arXiv for use of its open access interoperability."
        summary_header = "Summaries"
        if category == "CL":
            intro = (
                "If you would rather listen to today's summaries, you can hear them on the TechcraftingAI NLP podcast. ",
                "Your virtual host will be happy to read them to you!",
            )
        elif category == "RO":
            intro = (
                "If you would rather listen to today's summaries, you can hear them on the TechcraftingAI Robotics podcast. ",
                "Your virtual host will be happy to read them to you!",
            )

        intro_paragraph = doc.add_paragraph()
        intro_paragraph.add_run(intro)
        theme_header_heading = doc.add_heading(theme_header, level=2)
        theme_header_heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        theme_lines = themes.split("•")
        for line in theme_lines:
            if line == "":
                continue
            theme_paragraph = doc.add_paragraph()
            theme_paragraph.add_run("•" + line)
        thank_you_paragraph = doc.add_paragraph()
        thank_you_paragraph.add_run(thank_you)
        summary_header_heading = doc.add_heading(summary_header, level=2)
        summary_header_heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

        for record in records:
            if record["primary_category"] == category and record["date"] == research_date:
                title_pdf_paragraph = doc.add_paragraph()

                cleaned_title = re.sub(r"\n\s*", " ", record["title"])
                add_hyperlink(title_pdf_paragraph, cleaned_title, record["abstract_url"])

                title_pdf_paragraph.add_run(" [")
                add_hyperlink(title_pdf_paragraph, "PDF", record["abstract_url"].replace("abs", "pdf"))
                title_pdf_paragraph.add_run("]")

                authors = [f"{author['first_name']} {author['last_name']}" for author in record["authors"]]
                doc.add_paragraph("by " + ", ".join(authors))

                paragraphs = record["abstract"].split("\n\n")
                for p in paragraphs:
                    cleaned_paragraph = re.sub(r"\n\s*", " ", p)
                    no_latex_paragraph = latex_to_human_readable(cleaned_paragraph)
                    doc.add_paragraph(no_latex_paragraph)

        file_name = f"{research_date}_{group}_{category}_full_show_notes.docx"
        doc.save(os.path.join("show_notes", file_name))


def create_research_themes(research: str, category: str) -> str:
    global OPENAI_KEY
    system = "You are an expert machine learning researcher"
    user = ""
    instruction_start = "Give me the top five themes from the arXiv "
    category_text = ""
    instruction_end = "NLP research summaries below. These will be used in a social media post for a technical audience. Each theme must be a single short sentence. Be accurate, factual, do not exaggerate, and do not hallucinate. Readers will use themes to decide whether or not to read my newsletter. Your output must in json format and each theme must start with: •"
    if category == "CL":
        category_text = "NLP"
    elif category == "CV":
        category_text = "Computer Vision"
    elif category == "RO":
        category_text = "Robotics"

    user = f"{instruction_start}{category_text} {instruction_end}"

    client = OpenAI(api_key=OPENAI_KEY)
    response = client.chat.completions.create(
        model="gpt-4-1106-preview",
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user + "\n\n# Research Summaries\n\n" + research},
        ],
        temperature=0.9,
        max_tokens=1000,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0,
    )
    content = response.choices[0].message.content
    parsed_content = json.loads(content)
    themes = parsed_content.get("themes", [])
    return themes


def get_long_date(date: str):
    str_date = datetime.strptime(date, "%Y-%m-%d")
    long_date = str_date.strftime("%B %d, %Y")
    return long_date


def create_short_show_notes(categories: list, records: list, research_date: str, group: str):
    for category in categories:
        doc = Document()
        count = 0
        for record in records:
            if record["primary_category"] == category and record["date"] == date:
                count += 1
                title_pdf_paragraph = doc.add_paragraph()

                cleaned_title = re.sub(r"\n\s*", " ", record["title"])
                add_hyperlink(title_pdf_paragraph, cleaned_title, record["abstract_url"])

                title_pdf_paragraph.add_run(" [")
                add_hyperlink(title_pdf_paragraph, "PDF", record["abstract_url"].replace("abs", "pdf"))
                title_pdf_paragraph.add_run("]")

                authors = [f"{author['first_name']} {author['last_name']}" for author in record["authors"]]
                doc.add_paragraph("by " + ", ".join(authors))
                doc.add_paragraph()

        print(f"{category} {count}")


def create_script(categories, records, research_date, group):
    print(f"Creating script for {research_date}")
    print(f"Categories: {categories}")
    print(f"Records count: {len(records)}")
    for category in categories:
        doc = Document()
        # add intro notes
        intro = ""
        long_date = get_long_date(research_date)
        if category == "CL":
            intro = (
                "Hi, and welcome to Tech crafting AI NLP. I am your virtual host, Sage. "
                "Tech crafting AI NLP brings you daily summaries of new research released on archive. "
                "The podcast is produced by Brad Edwards, an AI Engineer from Vancouver, BC, and "
                "computer science graduate student at the University of York. Thank you to archive "
                "for use of its open access interoperability. \n\nHere is what was submitted on " + long_date + "."
            )
        elif category == "CV":
            intro = (
                "Hi, and welcome to Tech crafting AI Computer Vision. I am your virtual host, Sage. "
                "Tech crafting AI Computer Vision brings you daily summaries of new research released on archive. "
                "The podcast is produced by Brad Edwards, an AI Engineer from Vancouver, BC, and "
                "computer science graduate student at the University of York. Thank you to archive "
                "for use of its open access interoperability. \n\nHere is what was submitted on " + long_date + "."
            )
        elif category == "RO":
            intro = (
                "Hi, and welcome to Tech crafting AI Robotics. I am your virtual host, Sage. "
                "Tech crafting AI Robotics brings you daily summaries of new research released on archive. "
                "The podcast is produced by Brad Edwards, an AI Engineer from Vancouver, BC, and "
                "computer science graduate student at the University of York. Thank you to archive "
                "for use of its open access interoperability. \n\nHere is what was submitted on " + long_date + "."
            )

        intro_paragraph = doc.add_paragraph()
        intro_paragraph.add_run(intro)
        num_records = 0
        for record in records:
            if record["primary_category"] == category and record["date"] == research_date:
                title_pdf_paragraph = doc.add_paragraph()

                cleaned_title = re.sub(r"\n\s*", " ", record["title"])

                title_pdf_paragraph.add_run(cleaned_title)

                authors = [f"{author['first_name']} {author['last_name']}" for author in record["authors"]]
                doc.add_paragraph("by " + ", ".join(authors))
                doc.add_paragraph()
                paragraphs = record["abstract"].split("\n\n")
                for p in paragraphs:
                    cleaned_paragraph = re.sub(r"\n\s*", " ", p)
                    no_latex_paragraph = latex_to_human_readable(cleaned_paragraph)
                    doc.add_paragraph(no_latex_paragraph)

                doc.add_paragraph()
                num_records += 1

        if num_records == 0:
            continue
        outro = "That's all for today, thank you for listening. If you found the podcast helpful, please leave a comment, like, or share it with a friend. See you tomorrow!"
        outro_paragraph = doc.add_paragraph()
        outro_paragraph.add_run(outro)
        file_name = f"{research_date}_{group}_{category}_script.docx"
        doc.save(os.path.join("show_notes", file_name))
        themes = ""
        with open(os.path.join("show_notes", file_name), "rb") as f:
            document = Document(f)
            full_text = []
            for para in document.paragraphs:
                full_text.append(para.text)
            themes = create_research_themes("\n".join(full_text), category)
            print(f"Date: {research_date} Category: {category} Tokens: {len(full_text) / 6}")
        themes = "\n".join(themes)
        create_full_show_notes([category], records, research_date, "cs", themes)
        create_pod_notes([category], research_date, themes)
        create_post_text([category], research_date, themes)
        create_li_seo_lines(category, research_date)
        create_polly_audio("\n".join(full_text), file_name, research_date)


def create_pod_notes(categories: list, research_date: str, themes: str):
    for category in categories:
        category_text = ""
        if category == "CL":
            category_text = "NLP"
        elif category == "CV":
            category_text = "Computer Vision"
        elif category == "RO":
            category_text = "Robotics"

        text = f"arXiv {category_text} research summaries for {get_long_date(research_date)}.\n\nToday's Research Themes (AI-Generated):\n{themes}"
        file_name = f"{research_date}_{category}_pod_notes.txt"
        with open(os.path.join("show_notes", file_name), "w") as f:
            f.write(text)


def create_post_text(categories: list, research_date: str, themes: str):
    for category in categories:
        hashtags = ""
        if category == "CL":
            hashtags = "#naturallanguageprocessing #nlp #ai #artificialintelligence #llm"
        elif category == "CV":
            continue
        elif category == "RO":
            hashtags = "#robotics #ro #ai #artificialintelligence #llm"

        day_of_week = datetime.strptime(research_date, "%Y-%m-%d").strftime("%A")
        text = f"{day_of_week}'s Themes (AI-Generated):\n\n{themes}\n\nThank you to arXiv for use of its open access interoperability.\n\n{hashtags}"
        file_name = f"{research_date}_{category}_post_text.txt"
        with open(os.path.join("show_notes", file_name), "w") as f:
            f.write(text)


def create_li_seo_lines(category: str, research_date: str):
    long_date = get_long_date(research_date)
    text = ""
    description = ""
    if category == "CL":
        text = f"arXiv NLP research summaries for {long_date}."
        description = f"Research abstracts and links for arXiv NLP research summaries for {long_date}."
    elif category == "CV":
        text = f"arXiv Computer Vision research summaries for {long_date}."
        description = f"Research abstracts and links for arXiv Computer Vision research summaries for {long_date}."
    elif category == "RO":
        text = f"arXiv Robotics research summaries for {long_date}."
        description = f"Research abstracts and links for arXiv Robotics research summaries for {long_date}."

    file_name = f"{research_date}_{category}_li_seo_lines.txt"
    with open(os.path.join("show_notes", file_name), "w") as f:
        f.write(text)
        f.write("\n")
        f.write(description)


def insert_into_database(data: dict, db_config: dict) -> None:
    """Inserts the extracted fields into a database.

    Args:
        data (dict): A dictionary containing the extracted fields.
        db_config (dict): Database configuration details including host, user, password, etc.

    Returns:
        None
    """


# AtomikLabs code


def calculate_from_date() -> date:
    """Calculates from date for fetching summaries.

    Returns:
        date: From date.
    """
    today = datetime.today()
    yesterday = today - timedelta(days=0)
    return yesterday.date()


def log_initial_info(event: dict) -> None:
    """
    Logs initial info.

    Args:
        event (dict): Event.
    """
    logger.info(f"Received event: {event}")
    logger.info("Starting to fetch arXiv daily summaries")


def insert_fetch_status(date, aurora_cluster_arn, db_credentials_secret_arn, database):
    """
    Inserts fetch status as 'pending' for the given date using
    AWS RDSDataService.

    Args:
        date (date): Date for which to insert fetch status.
        aurora_cluster_arn (str): The ARN of the Aurora Serverless DB cluster.
        db_credentials_secret_arn (str): The ARN of the secret containing
                                         credentials to access the DB.
        database (str): Database name.
    """
    formatted_date = date.strftime("%Y-%m-%d")

    sql_statement = """
    INSERT INTO research_fetch_status (fetch_date, status)
    VALUES (CAST(:date AS DATE), 'pending') ON CONFLICT (fetch_date) DO NOTHING
    """

    parameters = [{"name": "date", "value": {"stringValue": formatted_date}}]

    response = execute_sql(sql_statement, parameters, aurora_cluster_arn, db_credentials_secret_arn, database)
    return response


def get_earliest_unfetched_date(aurora_cluster_arn, db_credentials_secret_arn, database, days=5) -> date:
    """
    Gets the earliest unfetched date using AWS RDSDataService.

    Args:
        aurora_cluster_arn (str): The ARN of the Aurora Serverless DB cluster.
        db_credentials_secret_arn (str): The ARN of the secret containing
                                         credentials to access the DB.
        database (str): Database name.
        days (int): Number of days to check for unfetched dates.

    Returns:
        date: Earliest unfetched date.
    """
    today = datetime.today().date()
    past_dates = [(today - timedelta(days=i)) for i in range(1, days + 1)]
    logger.info(f"Past dates: {past_dates}")
    logger.info(f"Today's date: {today}")

    placeholders = [f":date{i}" for i in range(len(past_dates))]
    placeholder_string = ", ".join(placeholders)

    sql_statement = f"""
    SELECT fetch_date FROM research_fetch_status
    WHERE fetch_date = ANY(ARRAY[{placeholder_string}]::DATE[]) AND status = 'success'
    """  # nosec

    parameters = [
        {"name": f"date{i}", "value": {"stringValue": date.strftime("%Y-%m-%d")}} for i, date in enumerate(past_dates)
    ]

    try:
        response = execute_sql(sql_statement, parameters, aurora_cluster_arn, db_credentials_secret_arn, database)

        fetched_dates = [
            datetime.strptime(result[0]["stringValue"], "%Y-%m-%d").date() for result in response["records"]
        ]
        unfetched_dates = list(set(past_dates) - set(fetched_dates))
        # prepend one day earlier than the earliest unfetched date (first date in the list)
        # arXiv doesn't always return the research for the date in the request (earliest date in this list)
        unfetched_dates.insert(0, unfetched_dates[0] - timedelta(days=1))
        logger.info(f"Unfetched dates: {unfetched_dates}")

        earliest_date = min(unfetched_dates) if unfetched_dates else None
    except Exception as e:
        logger.error(f"Database query failed: {str(e)}")
        earliest_date = None

    return earliest_date


def get_fetch_status(date: date, aurora_cluster_arn, db_credentials_secret_arn, database) -> str:
    """
    Gets fetch status for the given date using AWS RDSDataService.

    Args:
        date (date): Date for which to get fetch status.
        aurora_cluster_arn (str): The ARN of the Aurora Serverless DB cluster.
        db_credentials_secret_arn (str): The ARN of the secret containing
        credentials to access the DB.
        database (str): Database name.

    Returns:
        str: Fetch status.
    """
    formatted_date = date.strftime("%Y-%m-%d")

    sql_statement = """
    SELECT status FROM research_fetch_status
    WHERE fetch_date = CAST(:date AS DATE)
    """

    parameters = [{"name": "date", "value": {"stringValue": formatted_date}}]

    response = execute_sql(sql_statement, parameters, aurora_cluster_arn, db_credentials_secret_arn, database)
    logger.info(f"Fetch status response: {response} for date: {date}")
    if "records" in response and response["records"]:
        return response["records"][0][0].get("stringValue", "status_not_found")
    else:
        return "status_not_found"


def generate_date_list(start_date: date, end_date: date) -> List[date]:
    """
    Generates a list of dates between the given start and end dates.

    Args:
        start_date (date): Start date.
        end_date (date): End date.

    Returns:
        List[date]: List of dates.
    """
    delta = end_date - start_date
    if delta.days < 0:
        raise ValueError("End date must be after start date")

    dates = []
    for i in range(delta.days + 1):
        research_date = start_date + timedelta(days=i)
        fetch_status = get_fetch_status(research_date, AURORA_CLUSTER_ARN, DB_CREDENTIALS_SECRET_ARN, DATABASE)
        if fetch_status != "success":
            dates.append(research_date)
    return dates


def create_polly_audio(text: str, file_name: str, research_date: str) -> None:
    """
    Creates Polly audio for the given text.

    Args:
        text (str): Text to convert to audio.
        file_name (str): File name for the audio file.

    Returns:
        None
    """
    try:
        polly_client = boto3.client("polly")
        polly_client.start_speech_synthesis_task(
            OutputFormat="mp3",
            OutputS3BucketName=BUCKET_NAME,
            OutputS3KeyPrefix=f"pods/{research_date}-{file_name}",
            Text=text,
            TextType="text",
            VoiceId="Matthew",
        )
    except Exception as e:
        logger.error(f"Failed to create Polly audio for {file_name}")
        logger.error(e)


def schedule_for_later() -> None:
    """
    Schedules the Lambda function for later.
    """
    future_time = datetime.today() + timedelta(hours=5)

    cron_time = future_time.strftime("%M %H %d %m ? %Y")

    client = boto3.client("events")

    client.put_rule(Name="DynamicRule", ScheduleExpression=f"cron({cron_time})", State="ENABLED")

    lambda_arn = f"arn:aws:lambda:{os.environ['AWS_REGION']}:\
        {os.environ['AWS_ACCOUNT_ID']}:function:\
            {os.environ['AWS_LAMBDA_FUNCTION_NAME']}"

    client.put_targets(
        Rule="DynamicRule",
        Targets=[{"Id": "reschedule-{os.environ['AWS_LAMBDA_FUNCTION_NAME']}}", "Arn": lambda_arn}],
    )


def update_research_fetch_status(
    from_date: date,
    summary_set: str,
    bucket_name: str,
    aurora_cluster_arn: str,
    db_credentials_secret_arn: str,
    database: str,
    fetched_data: List[str],
) -> bool:
    """
    Checks if research was found for a given date and updates that
    date's research fetch status

    Args:
        from_date (date): Summary date.
        summary_set (str): Summary set.
        bucket_name (str): S3 bucket name.
        aurora_cluster_arn (str): The ARN of the Aurora Serverless DB cluster.
        db_credentials_secret_arn (str): The ARN of the secret containing
        credentials to access the DB.
        database (str): Database name.
        fetched_data (List[str]): List of XML responses.

    Returns:
        bool: True if fetch was successful, False otherwise.
    """
    pattern = r"<dc:description>.*?<dc:date>" + re.escape(from_date.strftime("%Y-%m-%d")) + r"</dc:date>"

    success = any(re.search(pattern, xml, re.DOTALL) for xml in fetched_data)

    if success:
        set_fetch_status(from_date, "success", aurora_cluster_arn, db_credentials_secret_arn, database)
    else:
        set_fetch_status(from_date, "failure", aurora_cluster_arn, db_credentials_secret_arn, database)

    return success


def set_fetch_status(date: date, status, aurora_cluster_arn, db_credentials_secret_arn, database):
    """
    Sets fetch status in the database using AWS RDSDataService.

    Args:
        date (date): Date for which to set fetch status.
        status (str): Status to set ('success' or 'failure').
        aurora_cluster_arn (str): The ARN of the Aurora Serverless DB cluster.
        db_credentials_secret_arn (str): The ARN of the secret containing
        credentials to access the DB.
        database (str): Database name.
    """
    try:
        sql_statement = "UPDATE research_fetch_status SET status = :status \
            WHERE fetch_date = CAST(:date AS DATE)"

        parameters = [
            {"name": "date", "value": {"stringValue": date.strftime("%Y-%m-%d")}},
            {"name": "status", "value": {"stringValue": status}},
        ]

        execute_sql(sql_statement, parameters, aurora_cluster_arn, db_credentials_secret_arn, database)
        return True
    except Exception as e:
        logger.error(f"Database query failed: {str(e)}")
        return False


def persist_authors(research_summary: dict, aurora_cluster_arn, db_credentials_secret_arn, database):
    """
    Persists authors in the database using AWS RDSDataService.

    Args:
        research_summary (str): Research summary in json.
        aurora_cluster_arn (str): The ARN of the Aurora Serverless DB cluster.
        db_credentials_secret_arn (str): The ARN of the secret containing
        credentials to access the DB.
        database (str): Database name.
    """
    try:
        author_ids = []
        sql_statement = "INSERT INTO research_authors (first_name, last_name) VALUES "
        values = []
        parameters = []
        for idx, author in enumerate(research_summary["authors"]):
            first_name_param = f"first_name_{idx}"
            last_name_param = f"last_name_{idx}"
            values.append(f"(:{first_name_param}, :{last_name_param})")
            parameters.extend(
                [
                    {"name": first_name_param, "value": {"stringValue": author["first_name"]}},
                    {"name": last_name_param, "value": {"stringValue": author["last_name"]}},
                ]
            )
        sql_statement += ", ".join(values) + " ON CONFLICT (first_name, last_name) DO NOTHING"
        response = execute_sql(sql_statement, parameters, aurora_cluster_arn, db_credentials_secret_arn, database)
        for author in research_summary["authors"]:
            sql_statement = (
                "SELECT author_id FROM research_authors WHERE first_name = :first_name AND last_name = :last_name"
            )
            parameters = [
                {"name": "first_name", "value": {"stringValue": author["first_name"]}},
                {"name": "last_name", "value": {"stringValue": author["last_name"]}},
            ]
            response = execute_sql(sql_statement, parameters, aurora_cluster_arn, db_credentials_secret_arn, database)
            author_ids.append(response["records"][0][0]["longValue"])
        return author_ids
    except Exception as e:
        logger.error(f"Database query failed: {str(e)}")


def persist_research_summaries(research_summaries, aurora_cluster_arn, db_credentials_secret_arn, database):
    records_added = 0
    authors_added = 0
    for summary in research_summaries:
        try:
            author_ids = persist_authors(summary, aurora_cluster_arn, db_credentials_secret_arn, database)
            authors_added += len(author_ids)
            primary_category_id = None
            if summary["primary_category"]:
                sql = "SELECT category_id FROM research_category WHERE name = :category_name"
                parameters = [{"name": "category_name", "value": {"stringValue": summary["primary_category"]}}]
                response = execute_sql(sql, parameters, aurora_cluster_arn, db_credentials_secret_arn, database)
                if response.get("records"):
                    primary_category_id = response["records"][0][0]["longValue"]

            sql = """
            INSERT INTO research (
                title,
                primary_category,
                summary,
                date,
                unique_identifier,
                abstract_url,
                full_text_url,
                stored_pdf_url,
                stored_full_text_url
            )
                VALUES (
                :title, :primary_category, :summary, CAST(:date AS DATE), :unique_identifier,
                :abstract_url, :full_text_url, :stored_pdf_url, :stored_full_text_url
            )
            RETURNING research_id
            """
            parameters = []
            if primary_category_id:
                parameters = [
                    {"name": "title", "value": {"stringValue": summary["title"]}},
                    {"name": "primary_category", "value": {"longValue": primary_category_id}},
                    {"name": "summary", "value": {"stringValue": summary["abstract"]}},
                    {"name": "date", "value": {"stringValue": summary["date"]}},
                    {"name": "unique_identifier", "value": {"stringValue": summary["identifier"]}},
                    {"name": "abstract_url", "value": {"stringValue": summary["abstract_url"]}},
                    {"name": "full_text_url", "value": {"stringValue": summary["abstract_url"].replace("abs", "pdf")}},
                    {"name": "stored_pdf_url", "value": {"stringValue": ""}},
                    {"name": "stored_full_text_url", "value": {"stringValue": ""}},
                ]
            else:
                parameters = [
                    {"name": "title", "value": {"stringValue": summary["title"]}},
                    {"name": "primary_category", "value": {"isNull": True}},
                    {"name": "summary", "value": {"stringValue": summary["abstract"]}},
                    {"name": "date", "value": {"stringValue": summary["date"]}},
                    {"name": "unique_identifier", "value": {"stringValue": summary["identifier"]}},
                    {"name": "abstract_url", "value": {"stringValue": summary["abstract_url"]}},
                    {"name": "full_text_url", "value": {"stringValue": summary["abstract_url"].replace("abs", "pdf")}},
                    {"name": "stored_pdf_url", "value": {"stringValue": ""}},
                    {"name": "stored_full_text_url", "value": {"stringValue": ""}},
                ]
            response = execute_sql(sql, parameters, aurora_cluster_arn, db_credentials_secret_arn, database)
            if response.get("records"):
                research_id = response["records"][0][0]["longValue"]
                records_added += 1
            else:
                print("No records were returned.")
            for author_id in author_ids:
                sql = """
                INSERT INTO research_author (research_id, author_id)
                VALUES (:research_id, :author_id)
                """
                parameters = [
                    {"name": "research_id", "value": {"longValue": research_id}},
                    {"name": "author_id", "value": {"longValue": author_id}},
                ]
                response = execute_sql(sql, parameters, aurora_cluster_arn, db_credentials_secret_arn, database)
            print("authors inserted")
            sql = """SELECT set_id FROM research_set WHERE name = :set_name"""
            parameters = [{"name": "set_name", "value": {"stringValue": summary["group"].upper()}}]
            response = execute_sql(sql, parameters, aurora_cluster_arn, db_credentials_secret_arn, database)
            if response.get("records"):
                set_id = response["records"][0][0]["longValue"]
                print(f"Set id: {set_id}")
                sql = """
                INSERT INTO research_set_research (set_id, research_id)
                VALUES (:set_id, :research_id)
                """
                parameters = [
                    {"name": "research_id", "value": {"longValue": research_id}},
                    {"name": "set_id", "value": {"longValue": set_id}},
                ]
                response = execute_sql(sql, parameters, aurora_cluster_arn, db_credentials_secret_arn, database)
            print("set inserted")
            category_ids = []
            for category in summary["categories"]:
                sql = """SELECT category_id FROM research_category WHERE name = :category_name"""
                parameters = [{"name": "category_name", "value": {"stringValue": category}}]
                response = execute_sql(sql, parameters, aurora_cluster_arn, db_credentials_secret_arn, database)
                if response.get("records"):
                    category_ids.append(response["records"][0][0]["longValue"])
                    print(f"Category id: {category_ids}")
            sql = """
            INSERT INTO research_category_research (category_id, research_id) VALUES (:category_id, :research_id)
            """
            for category_id in category_ids:
                parameters = [
                    {"name": "research_id", "value": {"longValue": research_id}},
                    {"name": "category_id", "value": {"longValue": category_id}},
                ]
                response = execute_sql(sql, parameters, aurora_cluster_arn, db_credentials_secret_arn, database)
        except Exception as e:
            logger.error(f"Database query failed: {str(e)}")


def execute_sql(
    sql_statement: str, parameters: List[dict], aurora_cluster_arn: str, db_credentials_secret_arn: str, database: str
) -> dict:
    """
    Executes the given SQL statement using AWS RDSDataService.

    Args:
        sql_statement (str): SQL statement to execute.
        parameters (List[dict]): List of parameters.
        aurora_cluster_arn (str): The ARN of the Aurora Serverless DB cluster.
        db_credentials_secret_arn (str): The ARN of the secret containing
        credentials to access the DB.
        database (str): Database name.

    Returns:
        dict: Response from RDSDataService.
    """
    client = boto3.client("rds-data")

    try:
        response = client.execute_statement(
            resourceArn=aurora_cluster_arn,
            secretArn=db_credentials_secret_arn,
            database=database,
            sql=sql_statement,
            parameters=parameters,
        )
        return response
    except Exception as e:
        logger.error(f"Database query failed: {str(e)}")
        return {}


def count_tokens(text: str) -> int:
    """
    Counts the number of tokens in the given text.

    Args:
        text (str): Text to count tokens for.

    Returns:
        int: Number of tokens.
    """
    return len(text) / 6


def lambda_handler(event: dict, context) -> dict:
    global AURORA_CLUSTER_ARN, BASE_URL, BUCKET_NAME, DB_CREDENTIALS_SECRET_ARN, DATABASE, SUMMARY_SET
    log_initial_info(event)

    today = calculate_from_date()

    logger.info(f"Today's date: {today}")
    AURORA_CLUSTER_ARN = os.environ.get("RESOURCE_ARN")
    BASE_URL = os.environ.get("BASE_URL")
    BUCKET_NAME = os.environ.get("BUCKET_NAME")
    DB_CREDENTIALS_SECRET_ARN = os.environ.get("SECRET_ARN")
    DATABASE = os.environ.get("DATABASE_NAME")
    SUMMARY_SET = os.environ.get("SUMMARY_SET")
    storage_manager = StorageManager(os.environ.get("S3_BUCKET_NAME"), logger)
    storage_manager.persist(os.environ.get("S3_STORAGE_KEY_PREFIX"), "Yep, worked")


def config_for_test():
    global AURORA_CLUSTER_ARN, BASE_URL, BUCKET_NAME, DB_CREDENTIALS_SECRET_ARN, DATABASE, SUMMARY_SET, OPENAI_KEY
    AURORA_CLUSTER_ARN = "arn:aws:rds:us-east-1:758145997264:cluster:atomiklabs-dev-aurora-cluster"
    BASE_URL = "http://export.arxiv.org/oai2"
    BUCKET_NAME = "dev-atomiklabs-data-bucket"
    DB_CREDENTIALS_SECRET_ARN = "arn:aws:secretsmanager:us-east-1:758145997264:secret:dev/database-credentials-TuF8OS"
    DATABASE = "atomiklabs_dev_database"
    SUMMARY_SET = "cs"
    load_dotenv()
    OPENAI_KEY = os.environ.get("OPENAI_KEY")


def run_aws_test():
    config_for_test()
    today = calculate_from_date()
    print(f"Today: {today}")
    log_initial_info({"test": "test"})
    insert_fetch_status(date.today(), AURORA_CLUSTER_ARN, DB_CREDENTIALS_SECRET_ARN, DATABASE)
    # earliest = get_earliest_unfetched_date(AURORA_CLUSTER_ARN, DB_CREDENTIALS_SECRET_ARN, DATABASE)
    earliest = datetime.today().date() - timedelta(days=5)
    print(f"Earliest: {earliest}")
    # date_list = generate_date_list(earliest, today)
    date_list = [datetime.today().date()]
    for i in range(1, 2):
        date_list.append(datetime.today().date() - timedelta(days=i))
    print(f"Date List: {date_list}")
    xml_data_list = fetch_data(BASE_URL, earliest, SUMMARY_SET)

    extracted_data = []
    for xml_data in xml_data_list:
        extracted_data.append(parse_xml_data(xml_data, earliest))
    print(len(extracted_data))
    if len(extracted_data) < 1:
        print("No data")
        return

    FILE_PATHS = {"records": "records.json"}

    write_to_files(extracted_data, FILE_PATHS)
    records = [record for data in extracted_data for record in data.get("records", [])]
    # for data in extracted_data:
    # persist_research_summaries(data['records'], AURORA_CLUSTER_ARN, DB_CREDENTIALS_SECRET_ARN, DATABASE)
    print(f"Records: {len(records)}")
    for research_date in date_list:
        r = research_date.strftime("%Y-%m-%d")
        create_script(["CL", "CV", "RO"], records, r, "cs")
