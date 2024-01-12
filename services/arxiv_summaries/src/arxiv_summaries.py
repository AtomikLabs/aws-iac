import requests
import json
import logging
import time
from html import unescape
from collections import defaultdict
from docx import Document
import os
import re
import docx
import boto3
import defusedxml.ElementTree as ET
from botocore.exceptions import NoRegionError
from datetime import datetime, timedelta, date
from typing import List
from anthropic import Anthropic, HUMAN_PROMPT, AI_PROMPT


logger = logging.getLogger(__name__)
logging.getLogger().setLevel(logging.INFO)
TEST = False
AURORA_CLUSTER_ARN = ''
BASE_URL = ''
BUCKET_NAME = ''
DB_CREDENTIALS_SECRET_ARN = ''
DATABASE = ''
SUMMARY_SET = ''
ANTHROPIC_KEY = ''

# Hardcoded dictionary for category lookup
cs_categories_inverted = {
    'Computer Science - Artifical Intelligence': 'AI',
    'Computer Science - Hardware Architecture': 'AR',
    'Computer Science - Computational Complexity': 'CC',
    'Computer Science - Computational Engineering, Finance, and Science': 'CE',
    'Computer Science - Computational Geometry': 'CG',
    'Computer Science - Computation and Language': 'CL',
    'Computer Science - Cryptography and Security': 'CR',
    'Computer Science - Computer Vision and Pattern Recognition': 'CV',
    'Computer Science - Computers and Society': 'CY',
    'Computer Science - Databases': 'DB',
    'Computer Science - Distributed, Parallel, and Cluster Computing': 'DC',
    'Computer Science - Digital Libraries': 'DL',
    'Computer Science - Discrete Mathematics': 'DM',
    'Computer Science - Data Structures and Algorithms': 'DS',
    'Computer Science - Emerging Technologies': 'ET',
    'Computer Science - Formal Languages and Automata Theory': 'FL',
    'Computer Science - General Literature': 'GL',
    'Computer Science - Graphics': 'GR',
    'Computer Science - Computer Science and Game Theory': 'GT',
    'Computer Science - Human-Computer Interaction': 'HC',
    'Computer Science - Information Retrieval': 'IR',
    'Computer Science - Information Theory': 'IT',
    'Computer Science - Machine Learning': 'LG',
    'Computer Science - Logic in Computer Science': 'LO',
    'Computer Science - Multiagent Systems': 'MA',
    'Computer Science - Multimedia': 'MM',
    'Computer Science - Mathematical Software': 'MS',
    'Computer Science - Numerical Analysis': 'NA',
    'Computer Science - Neural and Evolutionary Computing': 'NE',
    'Computer Science - Networking and Internet Architecture': 'NI',
    'Computer Science - Other Computer Science': 'OH',
    'Computer Science - Operating Systems': 'OS',
    'Computer Science - Performance': 'PF',
    'Computer Science - Programming Languages': 'PL',
    'Computer Science - Robotics': 'RO',
    'Computer Science - Symbolic Computation': 'SC',
    'Computer Science - Sound': 'SD',
    'Computer Science - Software Engineering': 'SE',
    'Computer Science - Social and Information Networks': 'SI',
    'Computer Science - Systems and Control': 'SY'
}


def fetch_data(base_url: str, from_date: str, set: str) -> list:
    full_xml_responses = []
    params = {
        'verb': 'ListRecords',
        'set': set,
        'metadataPrefix': 'oai_dc',
        'from': from_date
    }

    backoff_times = [30, 120]

    while True:
        try:
            logging.info(f"Fetching data with parameters: {params}")
            response = requests.get(base_url, params=params)
            response.raise_for_status()
            full_xml_responses.append(response.text)
            print(params)
            root = ET.fromstring(response.content)
            resumption_token_element = root.find(".//{http://www.openarchives.org/OAI/2.0/}resumptionToken")

            if resumption_token_element is not None and resumption_token_element.text:
                logging.info(f"Found resumptionToken: {resumption_token_element.text}")
                print(f"Found resumptionToken: {resumption_token_element.text}")
                time.sleep(5)
                params = {'verb': 'ListRecords', 'resumptionToken': resumption_token_element.text}
            else:
                break

        except requests.exceptions.HTTPError as e:
            logging.error(f"HTTP error occurred: {e}")
            logging.error(f"Response content: {response.text}")
            print(e)

            if response.status_code == 503:
                backoff_time = response.headers.get('Retry-After', backoff_times.pop(0) if backoff_times else 30)
                logging.warning(f"Received 503 error, backing off for {backoff_time} seconds.")
                print(f"Received 503 error, backing off for {backoff_time} seconds.")
                time.sleep(int(backoff_time))
                continue

            break

        except Exception as e:
            logging.error(f"An unexpected error occurred: {e}")
            break

    return full_xml_responses


def parse_xml_data(xml_data: str, from_date: str) -> dict:
    extracted_data_chunk = defaultdict(list)

    try:
        root = ET.fromstring(xml_data)
        ns = {
            'oai': 'http://www.openarchives.org/OAI/2.0/',
            'dc': 'http://purl.org/dc/elements/1.1/'
        }

        for record in root.findall(".//oai:record", ns):
            date_elements = record.findall(".//dc:date", ns)
            if len(date_elements) != 1:
                continue

            identifier = record.find(".//oai:identifier", ns).text
            abstract_url = record.find(".//dc:identifier", ns).text

            creators_elements = record.findall(".//dc:creator", ns)
            authors = []
            for creator in creators_elements:
                name_parts = creator.text.split(", ", 1)
                last_name = name_parts[0]
                first_name = name_parts[1] if len(name_parts) > 1 else ""
                authors.append({'last_name': last_name, 'first_name': first_name})

            # Find all subjects
            subjects_elements = record.findall(".//dc:subject", ns)
            categories = [cs_categories_inverted.get(subject.text, "") for subject in subjects_elements]

            # Primary category is the first one in the list
            primary_category = categories[0] if categories else ""

            abstract = record.find(".//dc:description", ns).text
            title = record.find(".//dc:title", ns).text
            date = date_elements[0].text
            group = 'cs'

            extracted_data_chunk['records'].append({
                'identifier': identifier,
                'abstract_url': abstract_url,
                'authors': authors,
                'primary_category': primary_category,
                'categories': categories,  # All categories
                'abstract': abstract,
                'title': title,
                'date': date,
                'group': group
            })

    except ET.ParseError as e:
        print(f"Parse error: {e}")

    return extracted_data_chunk


def write_to_files(data: list, file_paths: dict) -> None:
    """Writes the extracted fields to different files.

    Args:
        data (list): A list of dictionaries containing the extracted fields.
        file_paths (dict): A dictionary with field names as keys and file paths as values.

    Returns:
        None
    """
    try:
        # Aggregating all 'records' from each dictionary in the list
        aggregated_records = [record for d in data for record in d.get('records', [])]

        with open(file_paths['records'], 'w') as f:
            json_data = json.dumps(aggregated_records, indent=4)
            f.write(json_data)

        logging.info(f"Successfully wrote records data to {file_paths['records']}.")

    except Exception as e:
        logging.error(f"Failed to write data: {e}")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def latex_to_human_readable(latex_str):
    # Remove $...$ delimiters
    latex_str = re.sub(r'\$(.*?)\$', r'\1', latex_str)

    simple_latex_to_text = {
        '\\ll': '<<',
        '\\alpha': 'alpha',
        '\\epsilon': 'epsilon',
        '\\widetilde': 'widetilde',
        '\\in': 'in',
        '\\leq': '<=',
        '\\geq': '>=',
        '\\pm': '±',
        '\\times': 'x',
        '\\sim': '~',
        '\\approx': '≈',
        '\\neq': '≠',
        '\\cdot': '·',
        '\\ldots': '...',
        '\\cdots': '...',
        '\\vdots': '...',
        '\\ddots': '...',
        '\\forall': 'for all',
        '\\exists': 'exists',
        '\\nabla': 'nabla',
        '\\partial': 'partial',
        '\\{': '{',
        '\\}': '}',
        '\\:': ' ',  # Small space
        '\\,': ' ',  # Thin space
        '\\;': ' ',  # Thick space
        '\\!': '',  # Negative space
        '_': '_'    # Subscript
    }

    for latex, text in simple_latex_to_text.items():
        latex_str = latex_str.replace(latex, text)

    single_arg_pattern = re.compile(r'\\(\w+){(.*?)}')
    latex_str = single_arg_pattern.sub(r'\2', latex_str)

    latex_str = latex_str.replace("``", '"').replace("''", '"')

    latex_str = latex_str.replace("--", "–")

    return unescape(latex_str)


def create_full_show_notes(categories: list, records: list, research_date: str, group: str, themes: str):
    for category in categories:
        doc = Document()
        intro = ''
        theme_header = "Today's Themes (AI-Generated)"
        thank_you = "Thank you to arXiv for use of its open access interoperability."
        summary_header = "Summaries"
        if category == 'CL':
            intro = ("If you would rather listen to today's summaries, you can hear them on the TechcraftingAI NLP podcast. ",
                     "Your virtual host will be happy to read them to you!")

        intro_paragraph = doc.add_paragraph()
        intro_paragraph.add_run(intro)
        theme_header_heading = doc.add_heading(theme_header, level=2)
        theme_header_heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        theme_paragraph = doc.add_paragraph()
        theme_paragraph.add_run(themes)
        thank_you_paragraph = doc.add_paragraph()
        thank_you_paragraph.add_run(thank_you)
        summary_header_heading = doc.add_heading(summary_header, level=2)
        summary_header_heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

        for record in records:
            if record['primary_category'] == category and record['date'] == research_date:
                title_pdf_paragraph = doc.add_paragraph()

                cleaned_title = re.sub('\n\s*', ' ', record['title'])
                add_hyperlink(title_pdf_paragraph, cleaned_title, record['abstract_url'])

                title_pdf_paragraph.add_run(" [")
                add_hyperlink(title_pdf_paragraph, "PDF", record['abstract_url'].replace('abs', 'pdf'))
                title_pdf_paragraph.add_run("]")

                authors = [f"{author['first_name']} {author['last_name']}" for author in record['authors']]
                doc.add_paragraph('by ' + ', '.join(authors))

                paragraphs = record['abstract'].split('\n\n')
                for p in paragraphs:
                    cleaned_paragraph = re.sub('\n\s*', ' ', p)
                    no_latex_paragraph = latex_to_human_readable(cleaned_paragraph)
                    doc.add_paragraph(no_latex_paragraph)

        file_name = f"{research_date}_{group}_{category}_full_show_notes.docx"
        doc.save(os.path.join('show_notes', file_name))


def create_research_summary(research: str, category: str) -> str:
    anthropic = Anthropic(api_key=ANTHROPIC_KEY)
    instruction = ''
    if category == 'CL':
        instruction = ("# Instructions\nYou are an expert machine learning researcher. Give me the top five ",
                       "research themes from the arXiv NLP research summaries for my podcast below. ",
                       "Do not be flashy, be accurate, factual, and succinct. Themes should be a single sentence. ",
                       " The audience is technical. You must NOT respond with anything other than the themes in the ",
                       "format below. If you do, I will not be able to use your response. Use the following format:\n\n",
                       "• <theme 1>\n",
                       "• <theme 2>\n",
                       "• <theme 3>\n",
                       "• <theme 4>\n",
                       "• <theme 5>\n\n",
                       "# Research\n")
    elif category == 'CV':
        instruction = ("# Instructions\nYou are an expert machine learning researcher. Give me the top five ",
                       "Do not be flashy, be accurate, factual, and succinct. Themes should be a single sentence. ",
                       " The audience is technical. You must NOT respond with anything other than the themes in the ",
                       "format below. If you do, I will not be able to use your response. Use the following format:\n\n",
                       "• <theme 1>\n",
                       "• <theme 2>\n",
                       "• <theme 3>\n",
                       "• <theme 4>\n",
                       "• <theme 5>\n\n",
                       "# Research\n")
    elif category == 'RO':
        instruction = ("# Instructions\nYou are an expert machine learning researcher. Give me the top five ",
                       "Do not be flashy, be accurate, factual, and succinct. Themes should be a single sentence. ",
                       " The audience is technical. You must NOT respond with anything other than the themes in the ",
                       "format below. If you do, I will not be able to use your response. Use the following format:\n\n",
                       "• <theme 1>\n",
                       "• <theme 2>\n",
                       "• <theme 3>\n",
                       "• <theme 4>\n",
                       "• <theme 5>\n\n",
                       "# Research\n")

    prompt = f"{HUMAN_PROMPT} {instruction} Research: {research}{AI_PROMPT}"
    completion = anthropic.completions.create(
        model="claude-2.1",
        max_tokens_to_sample=4000,
        prompt=prompt
    )
    return completion.completion


def get_long_date(date: str):
    str_date = datetime.strptime(date, '%Y-%m-%d')
    long_date = str_date.strftime('%B %d, %Y')
    return long_date


def create_short_show_notes(categories: list, records: list, research_date: str, group: str):
    for category in categories:
        doc = Document()
        count = 0
        for record in records:
            if record['primary_category'] == category and record['date'] == date:
                count += 1
                title_pdf_paragraph = doc.add_paragraph()

                cleaned_title = re.sub('\n\s*', ' ', record['title'])
                add_hyperlink(title_pdf_paragraph, cleaned_title, record['abstract_url'])

                title_pdf_paragraph.add_run(" [")
                add_hyperlink(title_pdf_paragraph, "PDF", record['abstract_url'].replace('abs', 'pdf'))
                title_pdf_paragraph.add_run("]")

                authors = [f"{author['first_name']} {author['last_name']}" for author in record['authors']]
                doc.add_paragraph('by ' + ', '.join(authors))
                doc.add_paragraph()

        print(f"{category} {count}")


def create_script(categories, records, research_date, group):
    for category in categories:
        doc = Document()
        # add intro notes
        intro = ''
        long_date = get_long_date(research_date)
        if category == 'CL':
            intro = ("Hi, and welcome to Tech crafting AI NLP. I am your virtual host, Sage. "
                     "Tech crafting AI NLP brings you daily summaries of new research released on archive. "
                     "The podcast is produced by Brad Edwards, an AI Engineer from Vancouver, BC, and "
                     "computer science graduate student at the University of York. Thank you to archive "
                     "for use of its open access interoperability. \n\nHere is what was submitted on " + long_date + ".")
        elif category == 'CV':
            intro = ("Hi, and welcome to Tech crafting AI Computer Vision. I am your virtual host, Sage. "
                     "Tech crafting AI Computer Vision brings you daily summaries of new research released on archive. "
                     "The podcast is produced by Brad Edwards, an AI Engineer from Vancouver, BC, and "
                     "computer science graduate student at the University of York. Thank you to archive "
                     "for use of its open access interoperability. \n\nHere is what was submitted on " + long_date + ".")
        elif category == 'RO':
            intro = ("Hi, and welcome to Tech crafting AI Robotics. I am your virtual host, Sage. "
                     "Tech crafting AI Robotics brings you daily summaries of new research released on archive. "
                     "The podcast is produced by Brad Edwards, an AI Engineer from Vancouver, BC, and "
                     "computer science graduate student at the University of York. Thank you to archive "
                     "for use of its open access interoperability. \n\nHere is what was submitted on " + long_date + ".")

        intro_paragraph = doc.add_paragraph()
        intro_paragraph.add_run(intro)
        for record in records:
            if record['primary_category'] == category and record['date'] == research_date:
                title_pdf_paragraph = doc.add_paragraph()

                cleaned_title = re.sub('\n\s*', ' ', record['title'])

                title_pdf_paragraph.add_run(cleaned_title)
                
                authors = [f"{author['first_name']} {author['last_name']}" for author in record['authors']]
                doc.add_paragraph('by ' + ', '.join(authors))
                doc.add_paragraph()
                paragraphs = record['abstract'].split('\n\n')
                for p in paragraphs:
                    cleaned_paragraph = re.sub('\n\s*', ' ', p)
                    no_latex_paragraph = latex_to_human_readable(cleaned_paragraph)
                    doc.add_paragraph(no_latex_paragraph)

                doc.add_paragraph()

        outro = "That's all for today, thank you for listening. If you found the podcast helpful, please leave a comment, like, or share it with a friend. See you tomorrow!"
        outro_paragraph = doc.add_paragraph()
        outro_paragraph.add_run(outro)
        file_name = f"{research_date}_{group}_{category}_script.docx"
        doc.save(os.path.join('show_notes', file_name))
        themes = ''
        with open(os.path.join('show_notes', file_name), 'rb') as f:
            document = Document(f)
            full_text = []
            for para in document.paragraphs:
                full_text.append(para.text)
            themes = create_research_summary('\n'.join(full_text), category)
        lines = themes.split('\n')
        themes = [line for line in lines if line.startswith('•')]
        themes = '\n'.join(themes)
        create_full_show_notes([category], records, research_date, 'cs', themes)
        create_pod_notes([category], research_date, themes)
        create_post_text([category], research_date, themes)


def create_pod_notes(categories: list, research_date: str, themes: str):
    for category in categories:
        category_text = ''
        if category == 'CL':
            category_text = 'NLP'
        elif category == 'CV':
            category_text = 'Computer Vision'
        elif category == 'RO':
            category_text = 'Robotics'

        text = f"arXiv {category_text} research summaries for {get_long_date(research_date)}.\n\nToday's Research Themes (AI-Generated):\n\n{themes}"
        file_name = f"{research_date}_{category}_pod_notes.txt"
        with open(os.path.join('show_notes', file_name), 'w') as f:
            f.write(text)


def create_post_text(categories: list, research_date: str, themes: str):
    for category in categories:
        hashtags = ''
        if category == 'CL':
            hashtags = "#naturallanguageprocessing #nlp #ai #artificialintelligence #llm"
        elif category == 'CV':
            continue
        elif category == 'RO':
            hashtags = "#robotics #ro #ai #artificialintelligence #llm"

        day_of_week = datetime.strptime(research_date, '%Y-%m-%d').strftime('%A')
        text = f"{day_of_week}'s Themes (AI-Generated):\n\n{themes}\n\nThank you to arXiv for use of its open access interoperability.\n\n{hashtags}"
        file_name = f"{research_date}_{category}_post_text.txt"
        with open(os.path.join('show_notes', file_name), 'w') as f:
            f.write(text)


def insert_into_database(data: dict, db_config: dict) -> None:
    """Inserts the extracted fields into a database.

    Args:
        data (dict): A dictionary containing the extracted fields.
        db_config (dict): Database configuration details including host, user, password, etc.

    Returns:
        None
    """

# AtomikLabs code


def calculate_from_date() -> date:
    """Calculates from date for fetching summaries.

    Returns:
        date: From date.
    """
    today = datetime.today()
    yesterday = today - timedelta(days=0)
    return yesterday.date()


def log_initial_info(event: dict) -> None:
    """
    Logs initial info.

    Args:
        event (dict): Event.
    """
    logger.info(f"Received event: {event}")
    logger.info("Starting to fetch arXiv daily summaries")


def insert_fetch_status(date, aurora_cluster_arn, db_credentials_secret_arn, database):
    """
    Inserts fetch status as 'pending' for the given date using
    AWS RDSDataService.

    Args:
        date (date): Date for which to insert fetch status.
        aurora_cluster_arn (str): The ARN of the Aurora Serverless DB cluster.
        db_credentials_secret_arn (str): The ARN of the secret containing
                                         credentials to access the DB.
        database (str): Database name.
    """
    client = boto3.client("rds-data")
    formatted_date = date.strftime("%Y-%m-%d")

    sql_statement = """
    INSERT INTO research_fetch_status (fetch_date, status)
    VALUES (CAST(:date AS DATE), 'pending') ON CONFLICT (fetch_date) DO NOTHING
    """

    parameters = [{"name": "date", "value": {"stringValue": formatted_date}}]

    response = client.execute_statement(
        resourceArn=aurora_cluster_arn,
        secretArn=db_credentials_secret_arn,
        database=database,
        sql=sql_statement,
        parameters=parameters,
    )
    return response


def get_earliest_unfetched_date(aurora_cluster_arn, db_credentials_secret_arn, database, days=5) -> date:
    """
    Gets the earliest unfetched date using AWS RDSDataService.

    Args:
        aurora_cluster_arn (str): The ARN of the Aurora Serverless DB cluster.
        db_credentials_secret_arn (str): The ARN of the secret containing
                                         credentials to access the DB.
        database (str): Database name.
        days (int): Number of days to check for unfetched dates.

    Returns:
        date: Earliest unfetched date.
    """
    client = boto3.client("rds-data")
    today = datetime.today().date()
    past_dates = [(today - timedelta(days=i)) for i in range(1, days + 1)]
    logger.info(f"Past dates: {past_dates}")
    logger.info(f"Today's date: {today}")

    placeholders = [f":date{i}" for i in range(len(past_dates))]
    placeholder_string = ", ".join(placeholders)
    sql_statement = f"""
    SELECT fetch_date FROM research_fetch_status
    WHERE fetch_date = ANY(ARRAY[{placeholder_string}]::DATE[]) AND status = 'success'
    """

    parameters = [
        {"name": f"date{i}", "value": {"stringValue": date.strftime("%Y-%m-%d")}} for i, date in enumerate(past_dates)
    ]

    try:
        response = client.execute_statement(
            resourceArn=aurora_cluster_arn,
            secretArn=db_credentials_secret_arn,
            database=database,
            sql=sql_statement,
            parameters=parameters,
        )

        fetched_dates = [
            datetime.strptime(result[0]["stringValue"], "%Y-%m-%d").date() for result in response["records"]
        ]
        unfetched_dates = list(set(past_dates) - set(fetched_dates))
        logger.info(f"Unfetched dates: {unfetched_dates}")

        earliest_date = min(unfetched_dates) if unfetched_dates else None
    except Exception as e:
        logger.error(f"Database query failed: {str(e)}")
        earliest_date = None

    return earliest_date


def get_fetch_status(date: date, aurora_cluster_arn, db_credentials_secret_arn, database) -> str:
    """
    Gets fetch status for the given date using AWS RDSDataService.

    Args:
        date (date): Date for which to get fetch status.
        aurora_cluster_arn (str): The ARN of the Aurora Serverless DB cluster.
        db_credentials_secret_arn (str): The ARN of the secret containing
        credentials to access the DB.
        database (str): Database name.

    Returns:
        str: Fetch status.
    """
    client = boto3.client("rds-data")
    formatted_date = date.strftime("%Y-%m-%d")

    sql_statement = """
    SELECT status FROM research_fetch_status
    WHERE fetch_date = CAST(:date AS DATE)
    """

    parameters = [{"name": "date", "value": {"stringValue": formatted_date}}]

    response = client.execute_statement(
        resourceArn=aurora_cluster_arn,
        secretArn=db_credentials_secret_arn,
        database=database,
        sql=sql_statement,
        parameters=parameters,
    )
    logger.info(f"Fetch status response: {response} for date: {date}")
    if "records" in response and response["records"]:
        return response["records"][0][0].get("stringValue", "status_not_found")
    else:
        return "status_not_found"


def generate_date_list(start_date: date, end_date: date) -> List[date]:
    """
    Generates a list of dates between the given start and end dates.

    Args:
        start_date (date): Start date.
        end_date (date): End date.

    Returns:
        List[date]: List of dates.
    """
    delta = end_date - start_date
    if delta.days < 0:
        raise ValueError("End date must be after start date")
    
    dates = []
    for i in range(delta.days + 1):
        research_date = start_date + timedelta(days=i)
        fetch_status = get_fetch_status(research_date, AURORA_CLUSTER_ARN, DB_CREDENTIALS_SECRET_ARN, DATABASE)
        if fetch_status != 'success':
            dates.append(research_date)
    return dates


def lambda_handler(event: dict, context) -> dict:
    global AURORA_CLUSTER_ARN, BASE_URL, BUCKET_NAME, DB_CREDENTIALS_SECRET_ARN, DATABASE, SUMMARY_SET
    log_initial_info(event)

    today = calculate_from_date()

    logger.info(f"Today's date: {today}")
    AURORA_CLUSTER_ARN = os.environ.get('RESOURCE_ARN')
    BASE_URL = os.environ.get('BASE_URL')
    BUCKET_NAME = os.environ.get('BUCKET_NAME')
    DB_CREDENTIALS_SECRET_ARN = os.environ.get('SECRET_ARN')
    DATABASE = os.environ.get('DATABASE_NAME')
    SUMMARY_SET = os.environ.get('SUMMARY_SET')


def config_for_test():
    global AURORA_CLUSTER_ARN, BASE_URL, BUCKET_NAME, DB_CREDENTIALS_SECRET_ARN, DATABASE, SUMMARY_SET, ANTHROPIC_KEY
    AURORA_CLUSTER_ARN = "arn:aws:rds:us-east-1:758145997264:cluster:atomiklabs-dev-aurora-cluster"
    BASE_URL = "http://export.arxiv.org/oai2"
    BUCKET_NAME = "atomiklabs-data-bucket-dev"
    DB_CREDENTIALS_SECRET_ARN = "arn:aws:secretsmanager:us-east-1:758145997264:secret:dev/database-credentials-TuF8OS"
    DATABASE = "atomiklabs_dev_database"
    SUMMARY_SET = "cs"
    ANTHROPIC_KEY = "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"


def run_test():
    config_for_test()

    FROM_DATE = "2024-01-09"

    xml_data_list = fetch_data(BASE_URL, FROM_DATE, "cs")

    extracted_data = []

    for xml_data in xml_data_list:
        extracted_data.append(parse_xml_data(xml_data, FROM_DATE))

    print(len(extracted_data))
    print(extracted_data[1]['records'][0])

    FILE_PATHS = {
        'records': 'records.json'
    }

    write_to_files(extracted_data, FILE_PATHS)

    records = [record for data in extracted_data for record in data.get('records', [])]

    create_full_show_notes(['CL', 'CV', 'RO'], records, '2024-01-10', 'cs')
    create_script(['CL', 'CV', 'RO'], records, '2024-01-10', 'cs')


def run_aws_test():
    config_for_test()
    today = calculate_from_date()
    print(f'Today: {today}')
    log_initial_info({"test": "test"})
    # insert_fetch_status(date.today(), AURORA_CLUSTER_ARN, DB_CREDENTIALS_SECRET_ARN, DATABASE)
    earliest = get_earliest_unfetched_date(AURORA_CLUSTER_ARN, DB_CREDENTIALS_SECRET_ARN, DATABASE)
    print(f'Earliest: {earliest}')
    date_list = generate_date_list(earliest, today)
    print(f'Date List: {date_list}')
    xml_data_list = fetch_data(BASE_URL, earliest, SUMMARY_SET)
    extracted_data = []
    for xml_data in xml_data_list:
        extracted_data.append(parse_xml_data(xml_data, earliest))

    print(len(extracted_data))
    # print(extracted_data[1]['records'][0])

    FILE_PATHS = {
        'records': 'records.json'
    }

    write_to_files(extracted_data, FILE_PATHS)

    records = [record for data in extracted_data for record in data.get('records', [])]

    for research_date in date_list:
        r = research_date.strftime("%Y-%m-%d")
        create_script(['CL', 'CV', 'RO'], records, r, 'cs')        


if __name__ == '__main__':
    # run_test()
    run_aws_test()