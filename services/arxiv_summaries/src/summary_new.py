import json
import logging
import os
import re
import requests
import xml.etree.ElementTree as ET
import boto3
import time
from collections import defaultdict
from datetime import datetime
from html import unescape
from docx import Document
import docx.opc.constants

# Initialize logging
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# Constants
BASE_URL = "http://export.arxiv.org/oai2"
FROM_DATE = "2024-01-05"
BUCKET_NAME = os.environ.get("BUCKET_NAME")


def lambda_handler(event, context):
    try:
        xml_data_list = fetch_data(BASE_URL, FROM_DATE, "cs")
        extracted_data = [parse_xml_data(xml_data, FROM_DATE) for xml_data in xml_data_list]
        records = [record for data in extracted_data for record in data.get('records', [])]

        # Generate and upload documents to S3
        generate_and_upload_documents(['CL', 'CV', 'RO'], records, '2024-01-06', 'cs', 'full_show_notes')
        generate_and_upload_documents(['CL', 'CV', 'RO'], records, '2024-01-06', 'cs', 'script')
        generate_and_upload_documents(['CL', 'CV', 'RO'], records, '2024-01-07', 'cs', 'full_show_notes')
        generate_and_upload_documents(['CL', 'CV', 'RO'], records, '2024-01-07', 'cs', 'script')

        return {
            'statusCode': 200,
            'body': json.dumps('Processing and upload completed successfully.')
        }
    except Exception as e:
        logger.error(f"An error occurred: {e}")
        return {
            'statusCode': 500,
            'body': json.dumps('Error processing the request.')
        }


def fetch_data(base_url: str, from_date: str, set: str) -> list:
    full_xml_responses = []
    params = {
        'verb': 'ListRecords',
        'set': set,
        'metadataPrefix': 'oai_dc',
        'from': from_date
    }

    backoff_times = [30, 120]

    while True:
        try:
            logging.info(f"Fetching data with parameters: {params}")
            response = requests.get(base_url, params=params)
            response.raise_for_status()
            full_xml_responses.append(response.text)
            print(params)
            root = ET.fromstring(response.content)
            resumption_token_element = root.find(".//{http://www.openarchives.org/OAI/2.0/}resumptionToken")

            if resumption_token_element is not None and resumption_token_element.text:
                logging.info(f"Found resumptionToken: {resumption_token_element.text}")
                print(f"Found resumptionToken: {resumption_token_element.text}")
                time.sleep(5)
                params = {'verb': 'ListRecords', 'resumptionToken': resumption_token_element.text}
            else:
                break

        except requests.exceptions.HTTPError as e:
            logging.error(f"HTTP error occurred: {e}")
            logging.error(f"Response content: {response.text}")
            print(e)

            if response.status_code == 503:
                backoff_time = response.headers.get('Retry-After', backoff_times.pop(0) if backoff_times else 30)
                logging.warning(f"Received 503 error, backing off for {backoff_time} seconds.")
                print(f"Received 503 error, backing off for {backoff_time} seconds.")
                time.sleep(int(backoff_time))
                continue

            break

        except Exception as e:
            logging.error(f"An unexpected error occurred: {e}")
            break

    return full_xml_responses


def parse_xml_data(xml_data: str, from_date: str) -> dict:
    extracted_data_chunk = defaultdict(list)

    try:
        root = ET.fromstring(xml_data)
        ns = {
            'oai': 'http://www.openarchives.org/OAI/2.0/',
            'dc': 'http://purl.org/dc/elements/1.1/'
        }

        for record in root.findall(".//oai:record", ns):
            date_elements = record.findall(".//dc:date", ns)
            if len(date_elements) != 1:
                continue

            identifier = record.find(".//oai:identifier", ns).text
            abstract_url = record.find(".//dc:identifier", ns).text

            creators_elements = record.findall(".//dc:creator", ns)
            authors = []
            for creator in creators_elements:
                name_parts = creator.text.split(", ", 1)
                last_name = name_parts[0]
                first_name = name_parts[1] if len(name_parts) > 1 else ""
                authors.append({'last_name': last_name, 'first_name': first_name})

            # Find all subjects
            subjects_elements = record.findall(".//dc:subject", ns)
            categories = [cs_categories_inverted.get(subject.text, "") for subject in subjects_elements]

            # Primary category is the first one in the list
            primary_category = categories[0] if categories else ""

            abstract = record.find(".//dc:description", ns).text
            title = record.find(".//dc:title", ns).text
            date = date_elements[0].text
            group = 'cs'

            extracted_data_chunk['records'].append({
                'identifier': identifier,
                'abstract_url': abstract_url,
                'authors': authors,
                'primary_category': primary_category,
                'categories': categories,  # All categories
                'abstract': abstract,
                'title': title,
                'date': date,
                'group': group
            })

    except ET.ParseError as e:
        print(f"Parse error: {e}")

    return extracted_data_chunk


def write_to_files(data: list, file_paths: dict) -> None:
    """Writes the extracted fields to different files.

    Args:
        data (list): A list of dictionaries containing the extracted fields.
        file_paths (dict): A dictionary with field names as keys and file paths as values.

    Returns:
        None
    """
    try:
        # Aggregating all 'records' from each dictionary in the list
        aggregated_records = [record for d in data for record in d.get('records', [])]

        with open(file_paths['records'], 'w') as f:
            json_data = json.dumps(aggregated_records, indent=4)
            f.write(json_data)

        logging.info(f"Successfully wrote records data to {file_paths['records']}.")

    except Exception as e:
        logging.error(f"Failed to write data: {e}")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def latex_to_human_readable(latex_str):
    # Remove $...$ delimiters
    latex_str = re.sub(r'\$(.*?)\$', r'\1', latex_str)

    simple_latex_to_text = {
        '\\ll': '<<',
        '\\alpha': 'alpha',
        '\\epsilon': 'epsilon',
        '\\widetilde': 'widetilde',
        '\\in': 'in',
        '\\leq': '<=',
        '\\geq': '>=',
        '\\pm': '±',
        '\\times': 'x',
        '\\sim': '~',
        '\\approx': '≈',
        '\\neq': '≠',
        '\\cdot': '·',
        '\\ldots': '...',
        '\\cdots': '...',
        '\\vdots': '...',
        '\\ddots': '...',
        '\\forall': 'for all',
        '\\exists': 'exists',
        '\\nabla': 'nabla',
        '\\partial': 'partial',
        '\\{': '{',
        '\\}': '}',
        '\\:': ' ',  # Small space
        '\\,': ' ',  # Thin space
        '\\;': ' ',  # Thick space
        '\\!': '',  # Negative space
        '_': '_'    # Subscript
    }

    for latex, text in simple_latex_to_text.items():
        latex_str = latex_str.replace(latex, text)

    single_arg_pattern = re.compile(r'\\(\w+){(.*?)}')
    latex_str = single_arg_pattern.sub(r'\2', latex_str)

    latex_str = latex_str.replace("``", '"').replace("''", '"')

    latex_str = latex_str.replace("--", "–")

    return unescape(latex_str)


def create_full_show_notes(categories: list, records: list, date: str, group: str):
    for category in categories:
        doc = Document()
        # add intro paragraph gen

        for record in records:
            if record['primary_category'] == category and record['date'] == date:
                title_pdf_paragraph = doc.add_paragraph()

                cleaned_title = re.sub('\n\s*', ' ', record['title'])
                add_hyperlink(title_pdf_paragraph, cleaned_title, record['abstract_url'])

                title_pdf_paragraph.add_run(" [")
                add_hyperlink(title_pdf_paragraph, "PDF", record['abstract_url'].replace('abs', 'pdf'))
                title_pdf_paragraph.add_run("]")

                authors = [f"{author['first_name']} {author['last_name']}" for author in record['authors']]
                doc.add_paragraph('by ' + ', '.join(authors))

                paragraphs = record['abstract'].split('\n\n')
                for p in paragraphs:
                    cleaned_paragraph = re.sub('\n\s*', ' ', p)
                    no_latex_paragraph = latex_to_human_readable(cleaned_paragraph)
                    doc.add_paragraph(no_latex_paragraph)

        file_name = f"{date}_{group}_{category}_full_show_notes.docx"
        doc.save(os.path.join('show_notes', file_name))


def create_short_show_notes(categories: list, records: list, date: str, group: str):
    for category in categories:
        doc = Document()
        count = 0
        for record in records:
            if record['primary_category'] == category and record['date'] == date:
                count += 1
                title_pdf_paragraph = doc.add_paragraph()

                cleaned_title = re.sub('\n\s*', ' ', record['title'])
                add_hyperlink(title_pdf_paragraph, cleaned_title, record['abstract_url'])

                title_pdf_paragraph.add_run(" [")
                add_hyperlink(title_pdf_paragraph, "PDF", record['abstract_url'].replace('abs', 'pdf'))
                title_pdf_paragraph.add_run("]")

                authors = [f"{author['first_name']} {author['last_name']}" for author in record['authors']]
                doc.add_paragraph('by ' + ', '.join(authors))
                doc.add_paragraph()

        print(f"{category} {count}")


def create_script(categories, records, date, group):
    for category in categories:
        doc = Document()
        # add intro notes

        for record in records:
            if record['primary_category'] == category and record['date'] == date:
                title_pdf_paragraph = doc.add_paragraph()

                cleaned_title = re.sub('\n\s*', ' ', record['title'])

                title_pdf_paragraph.add_run(cleaned_title)

                authors = [f"{author['first_name']} {author['last_name']}" for author in record['authors']]
                doc.add_paragraph('by ' + ', '.join(authors))
                doc.add_paragraph()
                paragraphs = record['abstract'].split('\n\n')
                for p in paragraphs:
                    cleaned_paragraph = re.sub('\n\s*', ' ', p)
                    no_latex_paragraph = latex_to_human_readable(cleaned_paragraph)
                    doc.add_paragraph(no_latex_paragraph)

                doc.add_paragraph()

        file_name = f"{date}_{group}_{category}_script.docx"
        doc.save(os.path.join('show_notes', file_name))


def generate_and_upload_documents(categories, records, date, group, doc_type):
    if doc_type == 'full_show_notes':
        create_full_show_notes(categories, records, date, group)
    elif doc_type == 'script':
        create_script(categories, records, date, group)
    else:
        raise ValueError("Invalid document type")

    for category in categories:
        file_name = f"{date}_{group}_{category}_{doc_type}.docx"
        save_path = os.path.join('/tmp', file_name)
        upload_to_s3(save_path, BUCKET_NAME, file_name)


def test_fetch_and_process():
    xml_data_list = fetch_data(BASE_URL, FROM_DATE, "cs")
    extracted_data = [parse_xml_data(xml_data, FROM_DATE) for xml_data in xml_data_list]
    records = [record for data in extracted_data for record in data.get('records', [])]

    # Generate documents locally for testing
    generate_documents_locally(['CL', 'CV', 'RO'], records, '2024-01-05', 'cs', 'full_show_notes')
    generate_documents_locally(['CL', 'CV', 'RO'], records, '2024-01-04', 'cs', 'script')


def generate_documents_locally(categories, records, date, group, doc_type):
    if doc_type == 'full_show_notes':
        create_full_show_notes(categories, records, date, group)
    elif doc_type == 'script':
        create_script(categories, records, date, group)
    else:
        raise ValueError("Invalid document type")

    for category in categories:
        file_name = f"{date}_{group}_{category}_{doc_type}.docx"
        save_path = os.path.join('.', file_name)
        print(f"Document saved: {save_path}")


def upload_to_s3(file_path, bucket_name, object_name):
    """
    Upload a file to an S3 bucket.
    """
    s3_client = boto3.client('s3')
    try:
        response = s3_client.upload_file(file_path, bucket_name, object_name)
    except Exception as e:
        logger.error(f"Error uploading file to S3: {e}")
        raise e


cs_categories_inverted = {
    'Computer Science - Artifical Intelligence': 'AI',
    'Computer Science - Hardware Architecture': 'AR',
    'Computer Science - Computational Complexity': 'CC',
    'Computer Science - Computational Engineering, Finance, and Science': 'CE',
    'Computer Science - Computational Geometry': 'CG',
    'Computer Science - Computation and Language': 'CL',
    'Computer Science - Cryptography and Security': 'CR',
    'Computer Science - Computer Vision and Pattern Recognition': 'CV',
    'Computer Science - Computers and Society': 'CY',
    'Computer Science - Databases': 'DB',
    'Computer Science - Distributed, Parallel, and Cluster Computing': 'DC',
    'Computer Science - Digital Libraries': 'DL',
    'Computer Science - Discrete Mathematics': 'DM',
    'Computer Science - Data Structures and Algorithms': 'DS',
    'Computer Science - Emerging Technologies': 'ET',
    'Computer Science - Formal Languages and Automata Theory': 'FL',
    'Computer Science - General Literature': 'GL',
    'Computer Science - Graphics': 'GR',
    'Computer Science - Computer Science and Game Theory': 'GT',
    'Computer Science - Human-Computer Interaction': 'HC',
    'Computer Science - Information Retrieval': 'IR',
    'Computer Science - Information Theory': 'IT',
    'Computer Science - Machine Learning': 'LG',
    'Computer Science - Logic in Computer Science': 'LO',
    'Computer Science - Multiagent Systems': 'MA',
    'Computer Science - Multimedia': 'MM',
    'Computer Science - Mathematical Software': 'MS',
    'Computer Science - Numerical Analysis': 'NA',
    'Computer Science - Neural and Evolutionary Computing': 'NE',
    'Computer Science - Networking and Internet Architecture': 'NI',
    'Computer Science - Other Computer Science': 'OH',
    'Computer Science - Operating Systems': 'OS',
    'Computer Science - Performance': 'PF',
    'Computer Science - Programming Languages': 'PL',
    'Computer Science - Robotics': 'RO',
    'Computer Science - Symbolic Computation': 'SC',
    'Computer Science - Sound': 'SD',
    'Computer Science - Software Engineering': 'SE',
    'Computer Science - Social and Information Networks': 'SI',
    'Computer Science - Systems and Control': 'SY'
}


if __name__ == "__main__":
    test_fetch_and_process()